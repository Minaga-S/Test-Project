from pathlib import Path
import re
from datetime import date
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SOURCE_FILES = [
    Path('README.md'),
    Path('docs/README.md'),
    Path('docs/overview/system-overview.md'),
    Path('docs/tutorials/local-development.md'),
    Path('docs/tutorials/report-first-incident.md'),
    Path('docs/guides/architecture-and-request-flow.md'),
    Path('docs/guides/security-and-operations-gotchas.md'),
    Path('docs/manuals/api-reference.md'),
    Path('docs/manuals/data-model-reference.md'),
]

OUTPUT_PATH = Path('docs/Hotel-Cybersecurity-Technical-Documentation-Professional.docx')

SECTION_SUMMARIES = {
    'README.md': 'Simple explanation: This section introduces the project purpose and points to where each type of information lives.',
    'docs/README.md': 'Simple explanation: This is the documentation map, so readers can quickly choose the right path based on their role.',
    'docs/overview/system-overview.md': 'Simple explanation: This section explains what the system is, what it does, and how the main parts connect.',
    'docs/tutorials/local-development.md': 'Simple explanation: This tutorial helps a new developer run the platform locally and validate that it works.',
    'docs/tutorials/report-first-incident.md': 'Simple explanation: This tutorial walks through one real end-to-end workflow, from login to dashboard verification.',
    'docs/guides/architecture-and-request-flow.md': 'Simple explanation: This guide explains request flow and design choices so maintainers can reason about behavior and extensions.',
    'docs/guides/security-and-operations-gotchas.md': 'Simple explanation: This guide highlights non-obvious behavior and operational risks to avoid surprises in production.',
    'docs/manuals/api-reference.md': 'Simple explanation: This manual is the authoritative endpoint reference, including request expectations and status semantics.',
    'docs/manuals/data-model-reference.md': 'Simple explanation: This manual defines persistent entities, environment configuration, and defaults used by the system.',
}


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')

    text = OxmlElement('w:t')
    text.text = 'Right-click this table and choose "Update Field" after opening in Word.'

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(text)
    run._r.append(fld_char_end)


def add_cover_page(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('Hotel Cybersecurity Governance System')
    title_run.bold = True
    title_run.font.size = doc.styles['Title'].font.size

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Technical Documentation')
    subtitle_run.italic = True

    doc.add_paragraph('')
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('Document Type: Formal Technical Documentation\n')
    meta.add_run(f'Generated On: {date.today().isoformat()}\n')
    meta.add_run('Version: 1.0\n')
    meta.add_run('Prepared For: Project Stakeholders, Developers, QA, and Operations')

    doc.add_page_break()


def apply_heading(doc, text, level):
    level = max(1, min(level, 4))
    doc.add_heading(text.strip(), level=level)


def markdown_to_doc(doc, lines):
    in_code_block = False

    for raw_line in lines:
        line = raw_line.rstrip('\n')

        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Consolas'
            continue

        if not line.strip():
            doc.add_paragraph('')
            continue

        heading_match = re.match(r'^(#{1,6})\s+(.*)$', line)
        if heading_match:
            apply_heading(doc, heading_match.group(2), len(heading_match.group(1)) + 1)
            continue

        numbered_match = re.match(r'^\s*\d+\.\s+(.*)$', line)
        if numbered_match:
            doc.add_paragraph(numbered_match.group(1).strip(), style='List Number')
            continue

        bullet_match = re.match(r'^\s*[-*]\s+(.*)$', line)
        if bullet_match:
            doc.add_paragraph(bullet_match.group(1).strip(), style='List Bullet')
            continue

        doc.add_paragraph(line)


def main():
    doc = Document()

    add_cover_page(doc)

    doc.add_heading('Table of Contents', level=1)
    toc_paragraph = doc.add_paragraph()
    add_toc(toc_paragraph)
    doc.add_page_break()

    doc.add_heading('Document Overview', level=1)
    doc.add_paragraph(
        'This unified document compiles all project technical documentation into a single professional reference. '
        'It is structured to support different audiences: decision-makers, developers, testers, and operations teams.'
    )
    doc.add_paragraph(
        'Use tutorials for step-by-step execution, guides for understanding design and operational choices, '
        'and manuals for authoritative technical references.'
    )
    doc.add_page_break()

    for idx, path in enumerate(SOURCE_FILES):
        if not path.exists():
            continue

        doc.add_heading(f'Section Source: {path.as_posix()}', level=1)
        explanation = SECTION_SUMMARIES.get(path.as_posix(), 'Simple explanation: This section provides detailed project documentation content.')
        doc.add_paragraph(explanation)
        doc.add_paragraph('')

        lines = path.read_text(encoding='utf-8').splitlines()
        markdown_to_doc(doc, lines)

        if idx != len(SOURCE_FILES) - 1:
            doc.add_page_break()

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f'Created {OUTPUT_PATH.as_posix()}')


if __name__ == '__main__':
    main()

