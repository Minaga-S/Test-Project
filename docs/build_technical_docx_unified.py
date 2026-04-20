from datetime import date
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SOURCE_PATH = Path("docs/Hotel-Cybersecurity-Technical-Documentation-Professional.md")
OUTPUT_PATH = Path("docs/Hotel-Cybersecurity-Technical-Documentation-Unified-Standards.docx")
BACKEND_PATH = Path("backend")
FRONTEND_PATH = Path("frontend")

SECTION_TITLE_MAP = {
    "README.md": "System Purpose and Operational Context",
    "docs/README.md": "Documentation Governance and Navigation",
    "docs/overview/system-overview.md": "System Scope and High-Level Architecture",
    "docs/tutorials/local-development.md": "Local Development and Execution Standard",
    "docs/tutorials/report-first-incident.md": "Operational Workflow for First Incident Reporting",
    "docs/guides/architecture-and-request-flow.md": "Architecture and Request Flow",
    "docs/guides/security-and-operations-gotchas.md": "Security and Operational Constraints",
    "docs/manuals/api-reference.md": "API Contract and Integration Reference",
    "docs/manuals/data-model-reference.md": "Data Model and Configuration Reference",
}


def add_toc(paragraph):
    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = 'Right-click this table and choose "Update Field" after opening in Word.'

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(text)
    run._r.append(fld_char_end)


def add_cover_page(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Hotel Cybersecurity Governance System")
    title_run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Comprehensive Technical Documentation")
    subtitle_run.italic = True

    doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Document Type: Consolidated Standards-Style Technical Report\n")
    meta.add_run(f"Generated On: {date.today().isoformat()}\n")
    meta.add_run("Version: 1.0\n")
    meta.add_run("Prepared For: Engineering, QA, Operations, and Technical Leadership")

    doc.add_page_break()


def normalize_heading(text):
    cleaned = text.strip()

    # Convert source-referencing headings into report-style section headings.
    source_match = re.match(r"Section Source:\s*(.+)$", cleaned)
    if source_match:
        source_key = source_match.group(1).strip()
        return SECTION_TITLE_MAP.get(source_key, source_key)

    numbered_source_match = re.match(r"\d+\.\s+Section Source:\s*(.+)$", cleaned)
    if numbered_source_match:
        source_key = numbered_source_match.group(1).strip()
        return SECTION_TITLE_MAP.get(source_key, source_key)

    return cleaned


def add_justified_paragraph(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _is_markdown_reference_line(text):
    return bool(re.search(r"(^|[\s`/\-])[\w./\-]+\.md\b", text, flags=re.IGNORECASE))


def _is_source_toc_line(text):
    return bool(re.match(r"^\s*\d+\.\s+Section Source:\s+.+$", text))


def _collect_codebase_facts():
    route_files = sorted((BACKEND_PATH / "routes").glob("*.js")) if (BACKEND_PATH / "routes").exists() else []
    controller_files = sorted((BACKEND_PATH / "controllers").glob("*.js")) if (BACKEND_PATH / "controllers").exists() else []
    service_files = sorted((BACKEND_PATH / "services").glob("*.js")) if (BACKEND_PATH / "services").exists() else []
    model_files = sorted((BACKEND_PATH / "models").glob("*.js")) if (BACKEND_PATH / "models").exists() else []
    frontend_pages = sorted(FRONTEND_PATH.glob("*.html")) if FRONTEND_PATH.exists() else []
    frontend_modules = sorted((FRONTEND_PATH / "js").glob("*.js")) if (FRONTEND_PATH / "js").exists() else []

    endpoint_count = 0
    route_groups = []
    route_pattern = re.compile(r"router\.(get|post|put|patch|delete)\s*\(")
    mount_pattern = re.compile(r"app\.use\(\s*['\"](/api/[^'\"]*)['\"]")

    server_js = BACKEND_PATH / "server.js"
    mounted_paths = []
    if server_js.exists():
        server_text = server_js.read_text(encoding="utf-8", errors="ignore")
        mounted_paths = sorted(set(mount_pattern.findall(server_text)))

    for route_file in route_files:
        text = route_file.read_text(encoding="utf-8", errors="ignore")
        endpoint_count += len(route_pattern.findall(text))
        route_groups.append(route_file.stem)

    return {
        "route_files": route_files,
        "controller_files": controller_files,
        "service_files": service_files,
        "model_files": model_files,
        "frontend_pages": frontend_pages,
        "frontend_modules": frontend_modules,
        "endpoint_count": endpoint_count,
        "route_groups": route_groups,
        "mounted_paths": mounted_paths,
    }


def append_codebase_implementation_section(doc):
    facts = _collect_codebase_facts()

    doc.add_page_break()
    doc.add_heading("Current Implementation Baseline (Code-Derived)", level=1)

    add_justified_paragraph(
        doc,
        "This section is generated from the current codebase layout so the report reflects implementation state "
        "rather than only narrative source text. The backend presents a layered API structure with route, controller, "
        "service, and model boundaries, while the frontend remains a modular static web application with page-level "
        "entry points and shared JavaScript behavior.",
    )

    add_justified_paragraph(
        doc,
        f"The backend currently contains {len(facts['route_files'])} route modules, "
        f"{len(facts['controller_files'])} controller modules, {len(facts['service_files'])} service modules, "
        f"and {len(facts['model_files'])} persistence models. Static route inspection identifies "
        f"approximately {facts['endpoint_count']} endpoint declarations across the route layer.",
    )

    if facts["mounted_paths"]:
        mounted_paths = ", ".join(facts["mounted_paths"])
        add_justified_paragraph(
            doc,
            "Mounted API groups observed in the application bootstrap include "
            f"{mounted_paths}. These route groups indicate clear domain segmentation for authentication, asset and "
            "incident operations, analytical services, and audit/monitoring concerns.",
        )

    add_justified_paragraph(
        doc,
        f"The frontend currently includes {len(facts['frontend_pages'])} page templates and "
        f"{len(facts['frontend_modules'])} JavaScript modules in the shared client script layer. This supports "
        "separation between page rendering concerns and reusable transport/session logic.",
    )

    if facts["route_groups"]:
        route_groups = ", ".join(sorted(set(facts["route_groups"])))
        add_justified_paragraph(
            doc,
            "Primary backend route module groups detected in the codebase are "
            f"{route_groups}. This distribution confirms that the system maintains domain-driven routing organization "
            "suitable for continued scaling and targeted regression testing.",
        )


def render_markdown_to_docx(doc, lines):
    in_code_block = False
    code_language = ""
    code_lines = []
    is_skipping_source_toc = False

    for raw_line in lines:
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        if is_skipping_source_toc:
            if stripped == "---":
                is_skipping_source_toc = False
            continue

        code_fence_match = re.match(r"^```(.*)$", stripped)
        if code_fence_match:
            if not in_code_block:
                in_code_block = True
                code_language = code_fence_match.group(1).strip().lower()
                code_lines = []
            else:
                in_code_block = False

                if code_language == "mermaid":
                    doc.add_heading("Diagram Placeholder", level=3)
                    add_justified_paragraph(
                        doc,
                        "Insert the rendered diagram for the following Mermaid specification in this location "
                        "during final publishing.",
                    )
                    for mermaid_line in code_lines:
                        p = doc.add_paragraph()
                        run = p.add_run(mermaid_line)
                        run.font.name = "Consolas"
                else:
                    if code_language:
                        doc.add_paragraph(f"Code Snippet ({code_language})")
                    else:
                        doc.add_paragraph("Code Snippet")

                    for code_line in code_lines:
                        p = doc.add_paragraph()
                        run = p.add_run(code_line)
                        run.font.name = "Consolas"

                code_language = ""
                code_lines = []

            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not stripped:
            doc.add_paragraph("")
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            heading_label = heading_match.group(2).strip()
            if heading_label.lower() == "table of contents":
                is_skipping_source_toc = True
                continue

            level = min(4, max(1, len(heading_match.group(1))))
            heading_text = normalize_heading(heading_label)
            doc.add_heading(heading_text, level=level)
            continue

        if _is_source_toc_line(stripped):
            continue

        if _is_markdown_reference_line(stripped):
            continue

        # Render list lines as normal paragraphs to keep an academic narrative tone.
        bullet_match = re.match(r"^\s*[-*]\s+(.*)$", line)
        if bullet_match:
            add_justified_paragraph(doc, bullet_match.group(1).strip())
            continue

        numbered_match = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if numbered_match:
            add_justified_paragraph(doc, numbered_match.group(1).strip())
            continue

        add_justified_paragraph(doc, line)


def main():
    if not SOURCE_PATH.exists():
        raise FileNotFoundError(f"Missing source markdown: {SOURCE_PATH.as_posix()}")

    doc = Document()

    add_cover_page(doc)

    doc.add_heading("Table of Contents", level=1)
    toc_paragraph = doc.add_paragraph()
    add_toc(toc_paragraph)
    doc.add_page_break()

    doc.add_heading("Document Purpose", level=1)
    add_justified_paragraph(
        doc,
        "This document is a consolidated technical report intended to function as the single reference for "
        "the Hotel Cybersecurity Governance System. It preserves technical depth from the project materials "
        "while presenting the content in a professional report-oriented structure.",
    )
    doc.add_page_break()

    source_lines = SOURCE_PATH.read_text(encoding="utf-8").splitlines()
    render_markdown_to_docx(doc, source_lines)
    append_codebase_implementation_section(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)

    print(f"Created {OUTPUT_PATH.as_posix()}")


if __name__ == "__main__":
    main()
