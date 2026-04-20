from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle

SOURCE_MD = Path("technical-documentation/Hotel-Cybersecurity-Technical-Documentation-Standards.md")
OUTPUT_DOCX = Path("technical-documentation/Hotel-Cybersecurity-Technical-Documentation-Standards.docx")
OUTPUT_PDF = Path("technical-documentation/Hotel-Cybersecurity-Technical-Documentation-Standards.pdf")


def parse_markdown(lines):
    blocks = []
    in_code_block = False
    code_lines = []
    code_language = ""

    for raw_line in lines:
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        code_fence = re.match(r"^```(.*)$", stripped)
        if code_fence:
            if not in_code_block:
                in_code_block = True
                code_language = code_fence.group(1).strip()
                code_lines = []
            else:
                blocks.append({
                    "type": "code",
                    "language": code_language,
                    "content": "\n".join(code_lines),
                })
                in_code_block = False
                code_lines = []
                code_language = ""
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not stripped:
            blocks.append({"type": "blank"})
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            blocks.append({
                "type": "heading",
                "level": len(heading_match.group(1)),
                "text": heading_match.group(2).strip(),
            })
            continue

        bullet_match = re.match(r"^\s*[-*]\s+(.*)$", line)
        if bullet_match:
            blocks.append({"type": "bullet", "text": bullet_match.group(1).strip()})
            continue

        numbered_match = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if numbered_match:
            blocks.append({"type": "numbered", "text": numbered_match.group(1).strip()})
            continue

        if "|" in line and line.count("|") >= 2:
            columns = [col.strip() for col in line.strip().strip("|").split("|")]
            blocks.append({"type": "table_row", "columns": columns})
            continue

        blocks.append({"type": "paragraph", "text": line})

    return blocks


def render_docx(blocks):
    doc = Document()

    title = doc.add_paragraph("Hotel Cybersecurity Governance System Technical Documentation")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)

    subtitle = doc.add_paragraph("Standards-Style Unified Technical Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    subtitle.runs[0].font.size = Pt(11)

    doc.add_page_break()

    for block in blocks:
        block_type = block["type"]

        if block_type == "blank":
            doc.add_paragraph("")
            continue

        if block_type == "heading":
            level = min(4, max(1, block["level"]))
            doc.add_heading(block["text"], level=level)
            continue

        if block_type == "paragraph":
            p = doc.add_paragraph(block["text"])
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            continue

        if block_type == "bullet":
            p = doc.add_paragraph(block["text"], style="List Bullet")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            continue

        if block_type == "numbered":
            p = doc.add_paragraph(block["text"], style="List Number")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            continue

        if block_type == "code":
            if block["language"]:
                doc.add_paragraph(f"Code Snippet ({block['language']})")
            else:
                doc.add_paragraph("Code Snippet")

            for code_line in block["content"].splitlines() or [""]:
                code_p = doc.add_paragraph()
                run = code_p.add_run(code_line)
                run.font.name = "Consolas"
                run.font.size = Pt(10)
            continue

        if block_type == "table_row":
            doc.add_paragraph(" | ".join(block["columns"]))
            continue

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)


def render_pdf(blocks):
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "TitleStyle",
        parent=styles["Title"],
        fontSize=18,
        leading=22,
        alignment=1,
        spaceAfter=10,
    )
    subtitle_style = ParagraphStyle(
        "SubtitleStyle",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        alignment=1,
        spaceAfter=16,
    )
    heading_styles = {
        1: ParagraphStyle("H1", parent=styles["Heading1"], fontSize=15, leading=18, spaceBefore=10, spaceAfter=6),
        2: ParagraphStyle("H2", parent=styles["Heading2"], fontSize=13, leading=16, spaceBefore=8, spaceAfter=5),
        3: ParagraphStyle("H3", parent=styles["Heading3"], fontSize=11, leading=14, spaceBefore=6, spaceAfter=4),
        4: ParagraphStyle("H4", parent=styles["Heading4"], fontSize=10, leading=13, spaceBefore=5, spaceAfter=3),
    }
    body_style = ParagraphStyle("Body", parent=styles["BodyText"], fontSize=10, leading=14, alignment=4, spaceAfter=6)
    bullet_style = ParagraphStyle("BulletStyle", parent=body_style, leftIndent=14)
    code_style = ParagraphStyle("CodeStyle", parent=styles["Code"], fontName="Courier", fontSize=8.5, leading=11)

    story = []
    story.append(Paragraph("Hotel Cybersecurity Governance System Technical Documentation", title_style))
    story.append(Paragraph("Standards-Style Unified Technical Report", subtitle_style))
    story.append(Spacer(1, 6 * mm))

    pending_table = []

    def flush_table():
        nonlocal pending_table
        if not pending_table:
            return

        table = Table(pending_table, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e9eef5")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#9ca3af")),
            ("PADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(table)
        story.append(Spacer(1, 3 * mm))
        pending_table = []

    for block in blocks:
        block_type = block["type"]

        if block_type != "table_row":
            flush_table()

        if block_type == "blank":
            story.append(Spacer(1, 2 * mm))
            continue

        if block_type == "heading":
            level = min(4, max(1, block["level"]))
            story.append(Paragraph(block["text"], heading_styles[level]))
            continue

        if block_type == "paragraph":
            story.append(Paragraph(block["text"], body_style))
            continue

        if block_type == "bullet":
            story.append(Paragraph(f"• {block['text']}", bullet_style))
            continue

        if block_type == "numbered":
            story.append(Paragraph(block["text"], bullet_style))
            continue

        if block_type == "code":
            if block["language"]:
                story.append(Paragraph(f"Code Snippet ({block['language']})", heading_styles[4]))
            else:
                story.append(Paragraph("Code Snippet", heading_styles[4]))
            story.append(Preformatted(block["content"] or "", code_style))
            story.append(Spacer(1, 2 * mm))
            continue

        if block_type == "table_row":
            if set("".join(block["columns"])) <= set(":- "):
                continue
            pending_table.append(block["columns"])
            continue

    flush_table()

    OUTPUT_PDF.parent.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(
        str(OUTPUT_PDF),
        pagesize=A4,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
    )
    document.build(story)


def main():
    if not SOURCE_MD.exists():
        raise FileNotFoundError(f"Source markdown not found: {SOURCE_MD}")

    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    blocks = parse_markdown(lines)

    render_docx(blocks)
    render_pdf(blocks)

    print(f"Created {OUTPUT_DOCX.as_posix()}")
    print(f"Created {OUTPUT_PDF.as_posix()}")


if __name__ == "__main__":
    main()
